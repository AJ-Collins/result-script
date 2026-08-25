"""
TUM ICI - Special/Supplementary Results Word Doc Generator
=============================================================
Builds the summarized results Word document from the JSON produced by
extract.py, in the same layout as the manually-produced reference
(ICI_SUMMARIZED_RESULTS ... Final_senate.docx):

  - Letterhead block (TUM / ICI / title / series)
  - Main table: grouped by programme type (Master/Degree/Diploma/
    Certificate), then Year of Study (descending), each row split into
    SP (Special) / SUP (Supplementary) sub-columns for every metric
  - "Total no of students per School/Institute" summary table
  - Signature lines (Director's Sign / Chair of Senate)

IMPORTANT KNOWN LIMITATION (flagged, not silently guessed):
The source PDF marks Special ("#") vs Supplementary ("##") at the
individual UNIT level inside a student's row, not at the whole-student
level, and the marker can appear on a wrapped continuation line
disconnected from the rest of that student's row. That makes a reliable
per-student SP/SUP classification unsafe to infer automatically. Rather
than risk silently wrong SP/SUP numbers, every programme's combined
total (from the PDF's own SUMMARY blocks, which is unambiguous) is
placed in the SUP column and SP is left as "-". This mirrors the
approach previously used for this same limitation, and is called out
explicitly in the generated document and in this script.
"""

import sys
import json
import os
import re
import shutil
import zipfile
import tempfile
from collections import defaultdict
from datetime import date
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Book Antiqua"

def e(text):
    return xml_escape(str(text))

# ─────────────────────────────────────────────────────────────────────────────
# NESTED-AWARE TABLE BOUNDARY FINDER
# ─────────────────────────────────────────────────────────────────────────────

def find_table_end(content, tbl_start):
    depth = 0
    pos   = tbl_start
    while pos < len(content):
        no = content.find('<w:tbl>',  pos)
        nc = content.find('</w:tbl>', pos)
        if nc == -1:
            raise ValueError("Malformed XML: unmatched <w:tbl>")
        if no != -1 and no < nc:
            depth += 1
            pos = no + 7
        else:
            depth -= 1
            pos = nc + 8
            if depth == 0:
                return pos
    raise ValueError("Malformed XML: could not find matching </w:tbl>")

def find_all_top_level_tables(content):
    tables = []
    pos    = 0
    while True:
        start = content.find('<w:tbl>', pos)
        if start == -1:
            break
        end = find_table_end(content, start)
        tables.append((start, end))
        pos = end
    return tables


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def programme_type(name: str) -> str:
    n = name.upper()
    if n.startswith("MASTER"):
        return "Masters"
    if n.startswith("BACHELOR"):
        return "Degree"
    if n.startswith("DIPLOMA"):
        return "Diploma"
    if n.startswith("CERTIFICATE"):
        return "Certificate"
    return "Other"


TYPE_ORDER = ["Masters", "Degree", "Diploma", "Certificate", "Other"]


def title_case(name: str) -> str:
    return name.title()


def set_cell_text(cell, text, bold=False, size=9, align="center", vmerge=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(str(text))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, hex_color="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    # Set the grid (tblGrid) itself, not just cells, or Word can ignore it.
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gridCol, w in zip(tblGrid.findall(qn("w:gridCol")), widths_cm):
            gridCol.set(qn("w:w"), str(int(Cm(w).twips)))
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def set_table_borders(table):
    """Apply a thin single-line border to all sides of every cell."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


# ─────────────────────────────────────────────────────────────────────────────
# Aggregation
# ─────────────────────────────────────────────────────────────────────────────

def aggregate(records):
    """
    One row per (programme_name, year_of_study). If several intakes share the
    same programme+year (different intake codes), their pass/fail/absent/total
    are SUMMED into a single row, matching how the manual reference collapses
    intakes.
    """
    agg = defaultdict(lambda: {
        "registered": 0, "absent": 0, "sat_exam": 0,
        "pass": 0, "fail": 0, "total": 0, "classified": 0,
        "sp_registered": 0, "sp_absent": 0, "sp_sat_exam": 0,
        "sp_pass": 0, "sp_fail": 0,
        "sup_registered": 0, "sup_absent": 0, "sup_sat_exam": 0,
        "sup_pass": 0, "sup_fail": 0,
    })
    SP_SUP_KEYS = (
        "sp_registered", "sp_absent", "sp_sat_exam", "sp_pass", "sp_fail",
        "sup_registered", "sup_absent", "sup_sat_exam", "sup_pass", "sup_fail",
    )
    for r in records:
        key = (r["programme_name"], r["year_of_study"])
        a = agg[key]
        if r["summary_type"] in ("pass_fail", "competency"):
            a["registered"] += r["total"]
            a["absent"]     += r["absent"]
            a["sat_exam"]   += r["sat_exam"]
            a["pass"]       += r["pass"]
            a["fail"]       += r["fail"]
            a["total"]      += r["total"]
            # Only pass_fail records carry sp_*/sup_* fields; competency
            # sections have no '#'/'##' rows so they don't contribute.
            for k in SP_SUP_KEYS:
                a[k] += r.get(k, 0)
        elif r["summary_type"] == "classification":
            # Classification sheets: everyone listed there passed at some
            # level (the labels ARE pass grades), and register/sit/total
            # equal the classified count. No separate fail/absent line.
            a["registered"] += r["total"]
            a["sat_exam"]   += r["total"]
            a["pass"]       += r["total"]
            a["total"]      += r["total"]
            a["classified"] += r["total"]
    return agg


# ─────────────────────────────────────────────────────────────────────────────
# Table builders
# ─────────────────────────────────────────────────────────────────────────────

MAIN_METRICS = [
    ("No. Registered", "registered", "sp_registered", "sup_registered"),
    ("No. Absent",     "absent",     "sp_absent",     "sup_absent"),
    ("No. Sat Exam",   "sat_exam",   "sp_sat_exam",   "sup_sat_exam"),
    ("No. Pass",       "pass",       "sp_pass",       "sup_pass"),
    ("No. Fail",       "fail",       "sp_fail",       "sup_fail"),
]
EXTRA_SINGLE = ["Deregistered", "Disciplinary", "Incomplete", "Total"]


def build_main_table(doc, agg):
    n_metric_cols = len(MAIN_METRICS) * 2               # SP/SUP pairs
    n_cols = 1 + 1 + n_metric_cols + len(EXTRA_SINGLE)   # S/No + Name + metrics + extras
    table = doc.add_table(rows=0, cols=n_cols)
    set_table_borders(table)

    # Header row 1: group labels (blank over S/No, Name; metric name spans 2; extras span1)
    hdr1 = table.add_row().cells
    hdr2 = table.add_row().cells
    set_cell_text(hdr1[0], "S/No", bold=True)
    set_cell_text(hdr1[1], "Programme Name", bold=True)
    set_cell_text(hdr2[0], "", bold=True)
    set_cell_text(hdr2[1], "", bold=True)
    hdr1[0].merge(hdr2[0])
    hdr1[1].merge(hdr2[1])

    col = 2
    for label, _, _, _ in MAIN_METRICS:
        c1 = hdr1[col]
        c2 = hdr1[col + 1]
        merged = c1.merge(c2)
        set_cell_text(merged, label, bold=True)
        set_cell_text(hdr2[col], "SP", bold=True)
        set_cell_text(hdr2[col + 1], "SUP", bold=True)
        col += 2
    for label in EXTRA_SINGLE:
        c = hdr1[col]
        c2 = hdr1[col]  # single col, still spans the 2 header rows vertically
        merged = c.merge(hdr2[col])
        set_cell_text(merged, label, bold=True)
        col += 1

    for c in list(hdr1) + list(hdr2):
        shade_cell(c)

    # Grouped data rows
    by_type = defaultdict(lambda: defaultdict(list))
    for (prog, year), vals in agg.items():
        by_type[programme_type(prog)][year].append((prog, vals))

    serial = 1
    grand = defaultdict(int)
    for ptype in TYPE_ORDER:
        if ptype not in by_type:
            continue
        for year in sorted(by_type[ptype].keys(), reverse=True):
            # group header row
            grow = table.add_row().cells
            merged = grow[0]
            for c in grow[1:]:
                merged = merged.merge(c)
            set_cell_text(merged, f"{ptype}: Year of Study: {year}", bold=True, align="center")
            shade_cell(merged, "F2F2F2")

            for prog, vals in sorted(by_type[ptype][year], key=lambda x: x[0]):
                row = table.add_row().cells
                set_cell_text(row[0], serial)
                set_cell_text(row[1], title_case(prog), align="left")
                col = 2
                # Show '-' for SP/SUP if no SP/SUP data was computed for this
                # row (i.e. classification-only or competency-only programme).
                has_sp_sup = (vals["sp_registered"] + vals["sup_registered"]) > 0
                for _, total_key, sp_key, sup_key in MAIN_METRICS:
                    sp_val  = vals[sp_key]  if has_sp_sup else "-"
                    sup_val = vals[sup_key] if has_sp_sup else vals[total_key]
                    set_cell_text(row[col],     sp_val)
                    set_cell_text(row[col + 1], sup_val)
                    grand[total_key] += vals[total_key]
                    if has_sp_sup:
                        grand[sp_key]  += vals[sp_key]
                        grand[sup_key] += vals[sup_key]
                    col += 2
                set_cell_text(row[col], "-"); col += 1                      # Deregistered
                set_cell_text(row[col], "-"); col += 1                      # Disciplinary
                inc = max(0, vals["registered"] - vals["sat_exam"] - vals["absent"])
                set_cell_text(row[col], inc); col += 1                      # Incomplete
                set_cell_text(row[col], vals["total"]); col += 1            # Total
                grand["incomplete"] += inc
                grand["gtotal"] += vals["total"]
                serial += 1

    return table


def build_summary_table(doc, agg):
    prog_sum = defaultdict(lambda: {
        "registered": 0, "absent": 0, "sat_exam": 0,
        "pass": 0, "fail": 0, "total": 0, "classified": 0,
        "sp_registered": 0, "sp_absent": 0, "sp_sat_exam": 0,
        "sp_pass": 0, "sp_fail": 0,
        "sup_registered": 0, "sup_absent": 0, "sup_sat_exam": 0,
        "sup_pass": 0, "sup_fail": 0,
    })
    sum_keys = (
        "registered", "absent", "sat_exam", "pass", "fail", "total", "classified",
        "sp_registered", "sp_absent", "sp_sat_exam", "sp_pass", "sp_fail",
        "sup_registered", "sup_absent", "sup_sat_exam", "sup_pass", "sup_fail",
    )
    for (prog, year), vals in agg.items():
        p = prog_sum[prog]
        for k in sum_keys:
            p[k] += vals[k]

    n_metric_cols = len(MAIN_METRICS) * 2
    n_cols = 1 + 1 + n_metric_cols + 3 + 1  # S/No, Name, metrics, Deregistered/Disciplinary/Incomplete, Total, No.Classified
    # (Deregistered, Disciplinary, Incomplete, Total, No. Classified) = 5 single cols
    n_cols = 1 + 1 + n_metric_cols + 5

    table = doc.add_table(rows=0, cols=n_cols)
    set_table_borders(table)

    hdr1 = table.add_row().cells
    hdr2 = table.add_row().cells
    set_cell_text(hdr1[0], "S/No", bold=True); hdr1[0].merge(hdr2[0])
    set_cell_text(hdr1[1], "Programme Name", bold=True); hdr1[1].merge(hdr2[1])
    col = 2
    for label, _, _, _ in MAIN_METRICS:
        merged = hdr1[col].merge(hdr1[col + 1])
        set_cell_text(merged, label, bold=True)
        set_cell_text(hdr2[col], "SP", bold=True)
        set_cell_text(hdr2[col + 1], "SUP", bold=True)
        col += 2
    for label in ["Deregistered", "Disciplinary", "Incomplete", "Total", "No. Classified"]:
        merged = hdr1[col].merge(hdr2[col])
        set_cell_text(merged, label, bold=True)
        col += 1
    for c in list(hdr1) + list(hdr2):
        shade_cell(c)

    sno = 1
    totals = defaultdict(int)
    sorted_progs = sorted(prog_sum.keys(), key=lambda p: (TYPE_ORDER.index(programme_type(p)), p))
    for prog in sorted_progs:
        v = prog_sum[prog]
        row = table.add_row().cells
        set_cell_text(row[0], sno)
        set_cell_text(row[1], title_case(prog), align="left")
        col = 2
        has_sp_sup = (v["sp_registered"] + v["sup_registered"]) > 0
        for _, total_key, sp_key, sup_key in MAIN_METRICS:
            sp_val  = v[sp_key]  if has_sp_sup else "-"
            sup_val = v[sup_key] if has_sp_sup else v[total_key]
            set_cell_text(row[col],     sp_val)
            set_cell_text(row[col + 1], sup_val)
            totals[total_key] += v[total_key]
            if has_sp_sup:
                totals[sp_key]  += v[sp_key]
                totals[sup_key] += v[sup_key]
            col += 2
        set_cell_text(row[col], "-"); col += 1
        set_cell_text(row[col], "-"); col += 1
        inc = max(0, v["registered"] - v["sat_exam"] - v["absent"])
        set_cell_text(row[col], inc); col += 1
        set_cell_text(row[col], v["total"]); col += 1
        set_cell_text(row[col], v["classified"] if v["classified"] else "-"); col += 1
        totals["incomplete"] += inc
        totals["gtotal"]     += v["total"]
        totals["classified"] += v["classified"]
        sno += 1

    # TOTAL row
    row = table.add_row().cells
    merged = row[0].merge(row[1])
    set_cell_text(merged, "TOTAL", bold=True)
    col = 2
    for _, total_key, sp_key, sup_key in MAIN_METRICS:
        set_cell_text(row[col],     totals[sp_key],  bold=True)
        set_cell_text(row[col + 1], totals[sup_key], bold=True)
        col += 2
    set_cell_text(row[col], "-", bold=True); col += 1
    set_cell_text(row[col], "-", bold=True); col += 1
    set_cell_text(row[col], totals["incomplete"], bold=True); col += 1
    set_cell_text(row[col], totals["gtotal"], bold=True); col += 1
    set_cell_text(row[col], totals["classified"] if totals["classified"] else "-", bold=True); col += 1

    return table


# ─────────────────────────────────────────────────────────────────────────────
# Main generator
# ─────────────────────────────────────────────────────────────────────────────

def generate(template_path, json_path, output_path):
    with open(json_path) as f:
        data = json.load(f)
    records = data["records"]
    series  = data.get("series", "")
    ay      = data.get("academic_year", "")
    
    total_pages = data.get("total_pages", "?")
    gen_date = date.today().strftime("%d %B %Y").lstrip("0")

    agg = aggregate(records)

    work_dir   = tempfile.mkdtemp()
    unpack_dir = os.path.join(work_dir, "unpacked")
    with zipfile.ZipFile(template_path, "r") as z:
        z.extractall(unpack_dir)

    doc_xml_path = os.path.join(unpack_dir, "word", "document.xml")
    with open(doc_xml_path, encoding="utf-8") as f:
        content = f.read()

    # Replacements
    content = content.replace("{{ semister }}", "SPECIAL/SUPPLEMENTARY")
    content = content.replace("{{ academic_year }}", e(ay))
    content = content.replace("{{ series }}", e(series))
    content = content.replace("{{ total_pages }}", str(total_pages))
    content = content.replace("{{ generated_date }}", e(gen_date))

    tables = find_all_top_level_tables(content)
    assert len(tables) >= 3, f"Expected at least 3 tables in template, found {len(tables)}"

    t1_end    = tables[0][1]
    last_end  = tables[-1][1]

    chunk_between = content[t1_end : tables[1][0]]
    doc_c1_match  = re.search(
        r'<w:p[^>]*>(?:(?!</w:p>).)*DOC\.\s*C[0-9][^<]*(?:(?!</w:p>).)*</w:p>\s*',
        chunk_between, re.DOTALL
    )
    if doc_c1_match:
        cut_at = t1_end + doc_c1_match.start()
    else:
        cut_at = tables[1][0]

    before_data = content[:cut_at]
    after_data  = content[last_end:]

    MAGIC_MARKER = "MAGIC_INSERTION_MARKER"
    marker_xml = f'<w:p><w:r><w:t>{MAGIC_MARKER}</w:t></w:r></w:p>'
    new_content = before_data + marker_xml + after_data
    with open(doc_xml_path, "w", encoding="utf-8") as f:
        f.write(new_content)

    temp_doc_path = os.path.join(work_dir, "temp_supp.docx")
    with (zipfile.ZipFile(template_path, "r") as zin,
          zipfile.ZipFile(temp_doc_path, "w", zipfile.ZIP_DEFLATED) as zout):
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.write(doc_xml_path, "word/document.xml")
            else:
                zout.writestr(item, zin.read(item.filename))

    doc = Document(temp_doc_path)
    
    insert_before = None
    for p in doc.paragraphs:
        if MAGIC_MARKER in p.text:
            insert_before = p._p
            break
    
    if insert_before is None:
        insert_before = doc.paragraphs[0]._p if len(doc.paragraphs) > 0 else None

    def insert_element(element):
        if insert_before is not None:
            insert_before.addprevious(element)
        else:
            doc._body._body.append(element)

    p1 = doc.add_paragraph()
    insert_element(p1._p)
    
    # The SP/SUP note block has been removed as requested.

    spacer1 = doc.add_paragraph()
    insert_element(spacer1._p)

    t1 = build_main_table(doc, agg)
    set_col_widths(t1, [1.1, 5.2] + [1.15] * 10 + [1.7, 1.5, 1.5, 1.3])
    insert_element(t1._tbl)

    spacer2 = doc.add_paragraph()
    insert_element(spacer2._p)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Total no of students per School/Institute")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(9)
    insert_element(title_p._p)

    t2 = build_summary_table(doc, agg)
    set_col_widths(t2, [1.0, 5.0] + [1.1] * 10 + [1.5, 1.4, 1.4, 1.2, 1.4])
    insert_element(t2._tbl)

    spacer3 = doc.add_paragraph()
    insert_element(spacer3._p)

    if insert_before is None or insert_before.getparent() is None:
        pass
    else:
        insert_before.getparent().remove(insert_before)

    doc.save(output_path)
    shutil.rmtree(work_dir)
    print(f"Saved: {output_path}")
    print(f"Records used: {len(records)}  |  Programme-year rows: {len(agg)}")


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python generate.py <template.docx> <data.json> <output.docx>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2], sys.argv[3])
